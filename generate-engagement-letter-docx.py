#!/usr/bin/env python3
"""Generate Hillway-branded DOCX engagement letter + appendices for HPC/2026/CPR-001.

House style: Word built-in Title / Heading 1 / Heading 2 / Heading 3 / Normal / List Bullet,
matching the existing Hillway-Health-Safety-Policy.docx pattern. Logo top-right, compliance
footer with page numbers. Calibri throughout, body 11pt.

Reads three markdown source files in this directory and produces three DOCXs:

  1. HPC-2026-CPR-001-Engagement-Letter-DRAFT.md  -> HPC-2026-CPR-001-Engagement-Letter.docx
  2. HPC-2026-CPR-001-Appendix-B-DPA-DRAFT.md      -> HPC-2026-CPR-001-Appendix-B-DPA.docx
  3. HPC-2026-CPR-001-Appendix-C-Complaints-DRAFT.md -> HPC-2026-CPR-001-Appendix-C-Complaints.docx

The "DRAFTING NOTES" section at the foot of the engagement-letter source is stripped before render.

Usage: python3 generate-engagement-letter-docx.py
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO = Path(__file__).resolve().parent
LOGO = REPO / "logo.png"

# Brand colours
INK = RGBColor(26, 26, 26)
BODY = RGBColor(51, 51, 51)
MUTED = RGBColor(102, 102, 102)
ACCENT = RGBColor(189, 22, 34)

DOCS = [
    {
        "src": "HPC-2026-CPR-001-Engagement-Letter-DRAFT.md",
        "dst": "HPC-2026-CPR-001-Engagement-Letter.docx",
        "footer_ref": "HPC/2026/CPR-001 · Engagement Letter",
    },
    {
        "src": "HPC-2026-CPR-001-Appendix-B-DPA-DRAFT.md",
        "dst": "HPC-2026-CPR-001-Appendix-B-DPA.docx",
        "footer_ref": "HPC/2026/CPR-001 · Appendix B · Data Processing Agreement",
    },
    {
        "src": "HPC-2026-CPR-001-Appendix-C-Complaints-DRAFT.md",
        "dst": "HPC-2026-CPR-001-Appendix-C-Complaints.docx",
        "footer_ref": "HPC/2026/CPR-001 · Appendix C · Complaint Handling Procedure",
    },
]


# ---------- helpers ----------


def shade_cell(cell, fill_hex):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def add_horizontal_rule(paragraph):
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_left_border(paragraph):
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "BFBFBF")
    p_bdr.append(left)
    p_pr.append(p_bdr)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(fld_end)


# ---------- inline parser ----------

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text, base_bold=False, base_italic=False, base_size=Pt(11)):
    """Render markdown inline formatting: **bold**, *italic*, `code`, [text](url)."""
    if not text:
        return

    def make_run(
        content, bold=False, italic=False, mono=False, link=False, size=base_size
    ):
        r = paragraph.add_run(content)
        r.font.name = "Consolas" if mono else "Calibri"
        r.font.size = Pt(10) if mono else size
        r.font.color.rgb = ACCENT if link else (INK if bold else BODY)
        r.bold = bold or base_bold
        r.italic = italic or base_italic
        if link:
            r.font.underline = True

    def emit_segment(segment):
        pos = 0
        for m in INLINE_RE.finditer(segment):
            if m.start() > pos:
                make_run(segment[pos : m.start()])
            token = m.group(0)
            if token.startswith("**") and token.endswith("**"):
                make_run(token[2:-2], bold=True)
            elif token.startswith("*") and token.endswith("*"):
                make_run(token[1:-1], italic=True)
            elif token.startswith("`") and token.endswith("`"):
                make_run(token[1:-1], mono=True)
            elif token.startswith("[") and "](" in token:
                link_text = token[1 : token.index("](")]
                make_run(link_text, link=True)
            pos = m.end()
        if pos < len(segment):
            make_run(segment[pos:])

    # Split on <br> tags so cell content can be multi-line
    for i, part in enumerate(re.split(r"<br\s*/?>", text)):
        if i > 0:
            paragraph.add_run().add_break()
        emit_segment(part)


# ---------- document setup ----------


def setup_document(footer_ref):
    doc = Document()

    # Page setup — match H&S house style margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default Normal style — Calibri 11pt
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    # Title style tuning — keep large, centre, ink colour
    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(26)
    title.font.color.rgb = INK
    title.font.bold = True
    title.paragraph_format.space_after = Pt(4)

    # Heading 1 / 2 / 3 — Calibri, ink, bold
    for h_name, h_size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)):
        s = doc.styles[h_name]
        s.font.name = "Calibri"
        s.font.size = Pt(h_size)
        s.font.color.rgb = INK
        s.font.bold = True
        s.paragraph_format.space_before = Pt(14 if h_name == "Heading 1" else 10)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.keep_with_next = True

    # List Bullet — tighter spacing
    lb = doc.styles["List Bullet"]
    lb.font.name = "Calibri"
    lb.font.size = Pt(11)
    lb.font.color.rgb = BODY
    lb.paragraph_format.space_after = Pt(2)

    # Header — logo right
    header = doc.sections[0].header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO.exists():
        h_para.add_run().add_picture(str(LOGO), width=Inches(1.4))

    # Footer — compliance line + page number, centered
    footer = doc.sections[0].footer
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = f_para.add_run(f"Hillway · {footer_ref} · CONFIDENTIAL · Page ")
    f_run.font.name = "Calibri"
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = MUTED
    add_page_number_field(f_para)

    return doc


# ---------- block renderers ----------


def add_title(doc, text):
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.color.rgb = INK
    run.bold = True


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_inline(
        p, text, base_bold=True, base_size=Pt({1: 16, 2: 13, 3: 11}.get(level, 11))
    )


def add_paragraph(doc, text):
    p = doc.add_paragraph(style="Normal")
    add_inline(p, text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_inline(p, text)


def add_blockquote(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, text, base_italic=True)
    add_left_border(p)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_horizontal_rule(p)


def add_table_block(doc, header_row, rows):
    """Render a markdown table as a Word table.

    If the header row is empty (all cells blank), treat as a layout table:
    no grid, no header shading, no row banding — just clean side-by-side text.
    Used for signature blocks.
    """
    is_layout = all(c.strip() == "" for c in header_row)

    if is_layout:
        tbl = doc.add_table(rows=len(rows), cols=len(header_row))
        # No table style → invisible borders
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for ri, row in enumerate(rows):
            for ci, cell_text in enumerate(row):
                cell = tbl.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                add_inline(p, cell_text.strip())
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(4)
        return

    # Data table — header shading + row banding
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header_row))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row — white text on dark
    for ci, cell_text in enumerate(header_row):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        shade_cell(cell, "1A1A1A")
        p = cell.paragraphs[0]
        run = p.add_run(cell_text.strip())
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Body rows — alternating shade
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = tbl.rows[1 + ri].cells[ci]
            cell.text = ""
            if ri % 2 == 0:
                shade_cell(cell, "F2F2F2")
            p = cell.paragraphs[0]
            add_inline(p, cell_text.strip(), base_size=Pt(10))

    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


# ---------- markdown parser ----------

TABLE_LINE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_DIVIDER = re.compile(r"^\s*\|?\s*[:\-\|\s]+\|?\s*$")


def split_table_row(line):
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def parse_and_render(doc, src_text):
    """Walk through the markdown line-by-line; render each block."""
    lines = src_text.splitlines()
    i = 0
    title_done = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if not line.strip():
            i += 1
            continue

        # Drafting notes terminator — stop processing
        if line.strip().startswith("## DRAFTING NOTES"):
            break

        # Title (single H1 at top of document)
        if line.startswith("# ") and not title_done:
            add_title(doc, line[2:].strip())
            title_done = True
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            add_hr(doc)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            block = []
            while i < len(lines) and lines[i].startswith("> "):
                block.append(lines[i][2:].rstrip())
                i += 1
            add_blockquote(doc, " ".join(block))
            continue

        # Bullet list
        if line.lstrip().startswith("- ") or line.lstrip().startswith("* "):
            indent = len(line) - len(line.lstrip())
            text = line.lstrip()[2:]
            j = i + 1
            while (
                j < len(lines)
                and lines[j].startswith(" " * (indent + 2))
                and not lines[j].lstrip().startswith(("-", "*"))
            ):
                text += " " + lines[j].strip()
                j += 1
            add_bullet(doc, text)
            i = j
            continue

        # Tables
        if TABLE_LINE.match(line):
            header = split_table_row(line)
            i += 1
            if i < len(lines) and TABLE_DIVIDER.match(lines[i]):
                i += 1
            body = []
            while i < len(lines) and TABLE_LINE.match(lines[i]):
                body.append(split_table_row(lines[i]))
                i += 1
            add_table_block(doc, header, body)
            continue

        # Default — paragraph
        para_lines = [line]
        j = i + 1
        while (
            j < len(lines)
            and lines[j].strip()
            and not lines[j].lstrip().startswith(("#", "- ", "* ", "> ", "|"))
            and lines[j].strip() not in ("---", "***", "___")
        ):
            para_lines.append(lines[j].rstrip())
            j += 1
        add_paragraph(doc, " ".join(para_lines))
        i = j


# ---------- main ----------


def build_one(src_md, dst_docx, footer_ref):
    src = REPO / src_md
    dst = REPO / dst_docx
    if not src.exists():
        print(f"  SKIP {src_md}: source not found", file=sys.stderr)
        return False
    text = src.read_text(encoding="utf-8")
    doc = setup_document(footer_ref)
    parse_and_render(doc, text)
    doc.save(str(dst))
    size_kb = dst.stat().st_size // 1024
    print(f"  OK   {dst.name}  ({size_kb} KB)")
    return True


def main():
    print(f"Building Hillway-branded DOCX from markdown sources in {REPO}")
    if not LOGO.exists():
        print(f"  WARN: logo not found at {LOGO}", file=sys.stderr)
    ok = 0
    for d in DOCS:
        if build_one(d["src"], d["dst"], d["footer_ref"]):
            ok += 1
    print(f"\nDone — {ok}/{len(DOCS)} documents generated.")


if __name__ == "__main__":
    main()
